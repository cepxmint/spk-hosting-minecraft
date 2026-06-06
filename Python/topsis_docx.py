import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import io


def latex_to_image(latex_str, fname=None, dpi=150, fontsize=14):
    """Render LaTeX equation to an image (in-memory or file)."""
    fig, ax = plt.subplots(figsize=(6, 0.6))
    ax.axis("off")
    ax.text(0.5, 0.5, f"${latex_str}$", fontsize=fontsize,
            ha="center", va="center", transform=ax.transAxes)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                transparent=True, pad_inches=0.1)
    plt.close(fig)
    if fname:
        with open(fname, "wb") as f:
            f.write(buf.getvalue())
    buf.seek(0)
    return buf


def add_eq(doc, latex):
    """Add a paragraph containing a LaTeX equation rendered as image."""
    buf = latex_to_image(latex)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(4.5))
    return p

# ============================================================
# 1. BACA DATASET
# ============================================================
df = pd.read_csv("Dokumen/Dataset_Hosting_Minecraft.csv")

kriteria = ["Harga (USD/mo)", "RAM (GB)", "Storage (GB)", "Slot Pemain", "Uptime (%)", "vCPU"]
jenis = ["cost", "benefit", "benefit", "benefit", "benefit", "benefit"]  # jenis tiap kriteria

# Bobot (default equal — bisa diubah)
bobot = [0.20, 0.20, 0.10, 0.10, 0.25, 0.15]

# Alternatif
alt_names = df["Platform"] + " - " + df["Tier"]

X = df[kriteria].values.astype(float)
n_alt, n_krit = X.shape

# ============================================================
# 2. TOPSIS
# ============================================================

# --- Normalisasi ---
norm = np.sqrt((X ** 2).sum(axis=0))
R = X / norm

# --- Normalisasi terbobot ---
W = np.array(bobot)
Y = R * W

# --- Solusi ideal ---
A_pos = np.where(np.array(jenis) == "benefit", Y.max(axis=0), Y.min(axis=0))
A_neg = np.where(np.array(jenis) == "benefit", Y.min(axis=0), Y.max(axis=0))

# --- Jarak ---
D_pos = np.sqrt(((Y - A_pos) ** 2).sum(axis=1))
D_neg = np.sqrt(((Y - A_neg) ** 2).sum(axis=1))

# --- Preferensi ---
V = D_neg / (D_pos + D_neg)

# --- Ranking ---
ranking = np.argsort(-V)  # descending

df_result = df.copy()
df_result["Nilai Preferensi"] = V.round(6)
df_result["Ranking"] = 0
for rank, idx in enumerate(ranking, 1):
    df_result.loc[df.index[idx], "Ranking"] = rank
df_result = df_result.sort_values("Ranking")

print("[OK] TOPSIS selesai")
print(f"     Alternatif terbaik: {df_result.iloc[0]['Platform']} - {df_result.iloc[0]['Tier']} (V={df_result.iloc[0]['Nilai Preferensi']:.6f})")

# ============================================================
# 3. GENERATE DOCX
# ============================================================
doc = Document()

# --- Style defaults ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ---- COVER / TITLE ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LAPORAN PERHITUNGAN TOPSIS")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rekomendasi Platform Hosting Server Minecraft")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Jumlah alternatif: {n_alt} | Jumlah kriteria: {n_krit}")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()  # spacer

# ---- 1. DATA AWAL ----
doc.add_heading("1. Data Awal (Matriks Keputusan)", level=2)
p = doc.add_paragraph("Berikut adalah data dari setiap alternatif platform hosting server Minecraft:")

table = doc.add_table(rows=1 + len(df), cols=2 + n_krit)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "No"
hdr[1].text = "Alternatif"
for j, k in enumerate(kriteria):
    hdr[2 + j].text = k
for i, (_, row) in enumerate(df.iterrows(), 1):
    r = table.rows[i].cells
    r[0].text = str(i)
    r[1].text = f"{row['Platform']} - {row['Tier']}"
    for j, k in enumerate(kriteria):
        val = row[k]
        r[2 + j].text = f"{val:.2f}" if isinstance(val, float) else str(val)

doc.add_paragraph()

# ---- 2. JENIS & BOBOT ----
doc.add_heading("2. Jenis Kriteria dan Bobot", level=2)
table2 = doc.add_table(rows=2, cols=n_krit + 1)
table2.style = "Light Grid Accent 1"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = "Kriteria"
for j, k in enumerate(kriteria):
    hdr2[1 + j].text = k
r2 = table2.rows[1].cells
r2[0].text = "Jenis"
for j in range(n_krit):
    r2[1 + j].text = jenis[j].capitalize()

doc.add_paragraph()
doc.add_paragraph("Bobot yang digunakan:")
for j, k in enumerate(kriteria):
    doc.add_paragraph(f"  w{j+1} ({k}) = {bobot[j]}", style="List Bullet")

# ---- 3. NORMALISASI ----
doc.add_heading("3. Normalisasi Matriks", level=2)
doc.add_paragraph("Rumus normalisasi:")
add_eq(doc, r"r_{ij} = \frac{x_{ij}}{\sqrt{\sum_{j=1}^{n} x_{ij}^2}}")
doc.add_paragraph("diterapkan untuk setiap kolom kriteria.")

table3 = doc.add_table(rows=1 + n_alt, cols=2 + n_krit)
table3.style = "Light Grid Accent 1"
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
hdr3[0].text = "No"
hdr3[1].text = "Alternatif"
for j, k in enumerate(kriteria):
    hdr3[2 + j].text = k
for i in range(n_alt):
    r = table3.rows[1 + i].cells
    r[0].text = str(i + 1)
    r[1].text = alt_names.iloc[i] if hasattr(alt_names, 'iloc') else alt_names[i]
    for j in range(n_krit):
        r[2 + j].text = f"{R[i, j]:.6f}"

doc.add_paragraph()

# ---- 4. NORMALISASI TERBOBOT ----
doc.add_heading("4. Normalisasi Terbobot", level=2)
doc.add_paragraph("Rumus:")
add_eq(doc, r"y_{ij} = w_j \cdot r_{ij}")

table4 = doc.add_table(rows=1 + n_alt, cols=2 + n_krit)
table4.style = "Light Grid Accent 1"
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
hdr4[0].text = "No"
hdr4[1].text = "Alternatif"
for j, k in enumerate(kriteria):
    hdr4[2 + j].text = k
for i in range(n_alt):
    r = table4.rows[1 + i].cells
    r[0].text = str(i + 1)
    r[1].text = alt_names.iloc[i] if hasattr(alt_names, 'iloc') else alt_names[i]
    for j in range(n_krit):
        r[2 + j].text = f"{Y[i, j]:.6f}"

doc.add_paragraph()

# ---- 5. SOLUSI IDEAL ----
doc.add_heading("5. Solusi Ideal Positif dan Negatif", level=2)
doc.add_paragraph("Solusi ideal positif:")
add_eq(doc, r"A_j^+ = \{(\max y_{ij} \mid j \in \text{benefit}), (\min y_{ij} \mid j \in \text{cost})\}")
doc.add_paragraph("Solusi ideal negatif:")
add_eq(doc, r"A_j^- = \{(\min y_{ij} \mid j \in \text{benefit}), (\max y_{ij} \mid j \in \text{cost})\}")

table5 = doc.add_table(rows=3, cols=n_krit + 1)
table5.style = "Light Grid Accent 1"
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5 = table5.rows[0].cells
hdr5[0].text = "Solusi"
for j, k in enumerate(kriteria):
    hdr5[1 + j].text = k

r5a = table5.rows[1].cells
r5a[0].text = "Positif (A+)"
for j in range(n_krit):
    r5a[1 + j].text = f"{A_pos[j]:.6f}" if jenis[j] == "benefit" else f"{A_pos[j]:.6f} (cost)"

r5b = table5.rows[2].cells
r5b[0].text = "Negatif (A-)"
for j in range(n_krit):
    r5b[1 + j].text = f"{A_neg[j]:.6f}" if jenis[j] == "benefit" else f"{A_neg[j]:.6f} (cost)"

doc.add_paragraph()

# ---- 6. JARAK & PREFERENSI ----
doc.add_heading("6. Jarak dan Nilai Preferensi", level=2)
doc.add_paragraph("Rumus jarak positif:")
add_eq(doc, r"D_i^+ = \sqrt{\sum_{j=1}^{n} (y_{ij} - A_j^+)^2}")
doc.add_paragraph("Rumus jarak negatif:")
add_eq(doc, r"D_i^- = \sqrt{\sum_{j=1}^{n} (y_{ij} - A_j^-)^2}")
doc.add_paragraph("Rumus preferensi:")
add_eq(doc, r"V_i = \frac{D_i^-}{D_i^+ + D_i^-}")

table6 = doc.add_table(rows=1 + n_alt, cols=5)
table6.style = "Light Grid Accent 1"
table6.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr6 = table6.rows[0].cells
hdr6[0].text = "No"
hdr6[1].text = "Alternatif"
hdr6[2].text = "D+"
hdr6[3].text = "D-"
hdr6[4].text = "V (Preferensi)"
for i in range(n_alt):
    r = table6.rows[1 + i].cells
    r[0].text = str(i + 1)
    r[1].text = alt_names.iloc[i] if hasattr(alt_names, 'iloc') else alt_names[i]
    r[2].text = f"{D_pos[i]:.6f}"
    r[3].text = f"{D_neg[i]:.6f}"
    r[4].text = f"{V[i]:.6f}"

doc.add_paragraph()

# ---- 7. RANKING ----
doc.add_heading("7. Hasil Ranking", level=2)
p = doc.add_paragraph(
    "Semakin tinggi nilai preferensi (V), semakin baik alternatif tersebut. "
    "Berikut adalah urutan ranking dari yang terbaik hingga terendah:"
)

table7 = doc.add_table(rows=1 + n_alt, cols=4)
table7.style = "Light Grid Accent 1"
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr7 = table7.rows[0].cells
hdr7[0].text = "Ranking"
hdr7[1].text = "Alternatif"
hdr7[2].text = "Nilai V"
hdr7[3].text = "Keterangan"
for i, (_, row) in enumerate(df_result.iterrows()):
    r = table7.rows[1 + i].cells
    r[0].text = str(int(row["Ranking"]))
    r[1].text = f"{row['Platform']} - {row['Tier']}"
    r[2].text = f"{row['Nilai Preferensi']:.6f}"
    r[3].text = "TERBAIK" if i == 0 else ""

doc.add_paragraph()

# ---- 8. KESIMPULAN ----
doc.add_heading("8. Kesimpulan", level=2)
best = df_result.iloc[0]
p = doc.add_paragraph()
p.add_run(
    f"Berdasarkan perhitungan TOPSIS dengan bobot {dict(zip(kriteria, bobot))}, "
    f"alternatif terbaik adalah {best['Platform']} - {best['Tier']} "
    f"dengan nilai preferensi V = {best['Nilai Preferensi']:.6f}."
)

# Top 5
p = doc.add_paragraph("5 besar alternatif terbaik:")
top5 = df_result.head(5)
for i, (_, row) in enumerate(top5.iterrows(), 1):
    doc.add_paragraph(
        f"  {i}. {row['Platform']} - {row['Tier']} (V = {row['Nilai Preferensi']:.6f})",
        style="List Number"
    )

# --- SIMPAN ---
out_path = "Dokumen/Hasil_TOPSIS_Hosting_Minecraft.docx"
doc.save(out_path)
print(f"[OK] DOCX: {out_path}")
